import docx
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_report():
    doc = Document()
    
    # Page setup
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)
        
    # Styles helper
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)
    
    def set_cell_background(cell, fill_hex):
        tcPr = cell._tc.get_or_add_tcPr()
        shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
        tcPr.append(shd)

    def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
        tcPr = cell._tc.get_or_add_tcPr()
        tcMar = OxmlElement('w:tcMar')
        for margin, val in [('w:top', top), ('w:bottom', bottom), ('w:left', left), ('w:right', right)]:
            node = OxmlElement(margin)
            node.set(qn('w:w'), str(val))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tcPr.append(tcMar)

    def add_page_number(run):
        # python-docx doesn't support adding page numbers natively easily, 
        # so we will leave placeholders or standard footer spacing for MS Word.
        pass

    # Helper: Add Chapter Cover Page
    def add_chapter_cover(title_text):
        doc.add_page_break()
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        # Space before to center vertically
        p.paragraph_format.space_before = Pt(250)
        p.paragraph_format.space_after = Pt(20)
        run = p.add_run(title_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(36)
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
        
    # ==================== COVER PAGE ====================
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.paragraph_format.space_before = Pt(50)
    title_run = title_p.add_run("PACKMATE: A FULL-STACK TRAVEL PACKING LIST ORGANIZER\n")
    title_run.font.size = Pt(22)
    title_run.bold = True
    
    subtitle_p = doc.add_paragraph()
    subtitle_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle_p.add_run("A Project Report submitted in partial fulfillment of the requirements for the degree of Bachelor of Science / Computer Applications\n\n\n\n\n\n")
    sub_run.font.size = Pt(12)
    sub_run.italic = True
    
    details_p = doc.add_paragraph()
    details_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    d_run = details_p.add_run("Submitted By:\nYour Name (Reg. No: XXXX)\n\nUnder the Guidance of:\nGuide Name (Designation)\n\n\n\n\n\nDepartment of Computer Science\nCollege Name\n2025 - 2026")
    d_run.font.size = Pt(14)
    d_run.bold = True
    
    # ==================== CHAPTER 1 COVER ====================
    add_chapter_cover("INTRODUCTION")
    
    # ==================== CHAPTER 1 CONTENT ====================
    doc.add_page_break()
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. INTRODUCTION")
    h1_run.font.size = Pt(16)
    h1_run.bold = True
    h1.paragraph_format.space_after = Pt(12)
    
    h1_1 = doc.add_paragraph()
    h1_1_run = h1_1.add_run("1.1. SCOPE OF THE PROJECT")
    h1_1_run.font.size = Pt(14)
    h1_1_run.bold = True
    h1_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "This project focuses on developing a web-based travel packing list organizer system named \"PackMate\" "
        "to assist travelers in managing, organizing, and auditing their luggage items. It provides user "
        "registration and login features for Admin and Client roles. Clients can add, update, toggle packed status, "
        "and delete packing list items dynamically. The Admin role provides a centralized management platform to audit "
        "registered users and delete user accounts. The system ensures efficient data management through a structured "
        "NoSQL database, preventing last-minute packing omissions."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h1_2 = doc.add_paragraph()
    h1_2_run = h1_2.add_run("1.2. PROBLEM DESCRIPTION")
    h1_2_run.font.size = Pt(14)
    h1_2_run.bold = True
    h1_2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "In traditional travel planning, organizing a packing checklist is often a manual, paper-based, "
        "and error-prone process. Travelers frequently forget essential documents, medication, or chargers. "
        "Manual checks lead to delays, anxiety, and repetitive list preparation. Additionally, many travel checklist "
        "applications lack a responsive, beautiful web interface or fail to clean up database entities recursively "
        "when a user account is deleted. This project addresses these issues by providing a highly interactive, "
        "glassmorphic, purple-themed web platform for packing checklist CRUD and user database audits."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h1_2_1 = doc.add_paragraph()
    h1_2_1_run = h1_2_1.add_run("1.2.1. MODULE DESCRIPTION")
    h1_2_1_run.font.size = Pt(12)
    h1_2_1_run.bold = True
    h1_2_1.paragraph_format.space_after = Pt(6)
    
    modules = [
        "User Module",
        "Packing List Management Module",
        "Search and Filter Module",
        "Admin Module"
    ]
    for mod in modules:
        doc.add_paragraph(mod, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    # Detailed module text
    p = doc.add_paragraph()
    r = p.add_run("User Module\n")
    r.bold = True
    p.add_run("This module manages user registration and login for both Admin and Client roles. It ensures role-based secure login and dashboard redirection.")
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Packing List Management Module\n")
    r.bold = True
    p.add_run("This module allows logged-in clients to add new essential items, modify quantities, update item titles, and mark items as packed or unpacked dynamically.")
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Search and Filter Module\n")
    r.bold = True
    p.add_run("This module helps users search and filter their packing checklist items in real time as they type, allowing quick audit of packed vs unpacked items.")
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Admin Module\n")
    r.bold = True
    p.add_run("This module allows the administrator to view all registered user accounts and purge user data recursively. When a user is deleted, all their associated packing checklist items are cleared from the database automatically.")
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 2 COVER ====================
    add_chapter_cover("SYSTEM STUDY AND ANALYSIS")
    
    # ==================== CHAPTER 2 CONTENT ====================
    doc.add_page_break()
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. SYSTEM STUDY AND ANALYSIS")
    h2_run.font.size = Pt(16)
    h2_run.bold = True
    h2.paragraph_format.space_after = Pt(12)
    
    h2_1 = doc.add_paragraph()
    h2_1_run = h2_1.add_run("2.1. EXISTING SYSTEM")
    h2_1_run.font.size = Pt(14)
    h2_1_run.bold = True
    h2_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "The existing system for managing packing lists relies on paper checklists or basic offline notes applications. "
        "These traditional methods are highly inefficient as lists cannot be shared, dynamic progress calculations "
        "are not computed, and items cannot be easily updated. Information regarding travel guidelines and essentials "
        "remains unorganized. There is no role-based separation of data and no administrative control to monitor active "
        "user bases. Overall, the existing system is static, error-prone, and lacks user-friendliness."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h2_2 = doc.add_paragraph()
    h2_2_run = h2_2.add_run("2.2. PROPOSED SYSTEM")
    h2_2_run.font.size = Pt(14)
    h2_2_run.bold = True
    h2_2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "The proposed system is a full-stack web-based travel organizer platform. It provides a secure role-based login "
        "redirecting Users and Admins to separate dashboards. Users can execute CRUD actions to modify their list and "
        "view real-time recalculated statistics (total, packed, remaining items, progress bar). Administrators can audit "
        "all user credentials and delete accounts, triggering a database cascade delete that removes all associated travel lists. "
        "The layout utilizes premium dark CSS glassmorphism for enhanced accessibility and visual excellence."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h2_3 = doc.add_paragraph()
    h2_3_run = h2_3.add_run("2.3. FEASIBILITY STUDY")
    h2_3_run.font.size = Pt(14)
    h2_3_run.bold = True
    h2_3.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "The feasibility study evaluates if the proposed travel organizer web application is practical, cost-effective, "
        "and technically achievable. This analysis ensures the implementation runs smoothly under current resource structures."
    )
    p.paragraph_format.space_after = Pt(6)
    
    feasibility_points = [
        "Technical Feasibility",
        "Economic Feasibility",
        "Operational Feasibility"
    ]
    for pt in feasibility_points:
        doc.add_paragraph(pt, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    r = p.add_run("Technical Feasibility\n")
    r.bold = True
    p.add_run(
        "Technical feasibility examines if the required web technologies are accessible. The system is built "
        "using standard web technologies: HTML, CSS, JavaScript (React JS, Vite) for the frontend, and Java "
        "(Spring Boot, Maven) for the backend, with MongoDB Atlas as the NoSQL database. These tools are open-source, "
        "highly stable, and support clean local development and cloud hosting pipelines."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Economic Feasibility\n")
    r.bold = True
    p.add_run(
        "Economic feasibility determines if the project is financially viable. PackMate uses free, open-source software "
        "and tools, eliminating expensive licensing costs. Cloud deployments are hosted on free tiers: Vercel for "
        "static React compile, Render for Spring Boot hosting, and MongoDB Atlas for database clusters. Therefore, "
        "the cost of development and maintenance is extremely minimal, making the project highly viable."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Operational Feasibility\n")
    r.bold = True
    p.add_run(
        "Operational feasibility assesses how well the web system is accepted and operated by users. PackMate is "
        "designed with an intuitive, glassmorphic dashboard containing explicit progress bars, toggle controls, and "
        "modal forms. The Admin Dashboard is simplified to tables with delete warnings. No advanced technical training "
        "is required, ensuring high operational acceptance."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 3 COVER ====================
    add_chapter_cover("DEVELOPMENT ENVIRONMENT")
    
    # ==================== CHAPTER 3 CONTENT ====================
    doc.add_page_break()
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. DEVELOPMENT ENVIRONMENT")
    h3_run.font.size = Pt(16)
    h3_run.bold = True
    h3.paragraph_format.space_after = Pt(12)
    
    h3_1 = doc.add_paragraph()
    h3_1_run = h3_1.add_run("3.1. HARDWARE REQUIREMENTS")
    h3_1_run.font.size = Pt(14)
    h3_1_run.bold = True
    h3_1.paragraph_format.space_after = Pt(6)
    
    # Hardware Table
    hw_table = doc.add_table(rows=6, cols=2)
    hw_table.style = 'Light Shading Accent 1'
    hw_headers = ["Resource Spec", "Details"]
    hw_rows = [
        ("Processor", "Intel Core i5 / AMD Ryzen 5 or Above"),
        ("Processor Speed", "2.0 GHz or Above"),
        ("Ram", "8 GB or Above"),
        ("Hard Disk capacity", "2 GB or Above"),
        ("Monitor", "14 inches or Above")
    ]
    
    # Write Headers
    hdr_cells = hw_table.rows[0].cells
    hdr_cells[0].text = hw_headers[0]
    hdr_cells[1].text = hw_headers[1]
    for cell in hdr_cells:
        set_cell_background(cell, "7C4DFF")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell)
        
    # Write Rows
    for idx, (spec, det) in enumerate(hw_rows):
        row_cells = hw_table.rows[idx+1].cells
        row_cells[0].text = spec
        row_cells[1].text = det
        for cell in row_cells:
            set_cell_background(cell, "F3E5F5" if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    h3_2 = doc.add_paragraph()
    h3_2_run = h3_2.add_run("3.2. SOFTWARE SPECIFICATION")
    h3_2_run.font.size = Pt(14)
    h3_2_run.bold = True
    h3_2.paragraph_format.space_after = Pt(6)
    
    # Software Table
    sw_table = doc.add_table(rows=6, cols=2)
    sw_table.style = 'Light Shading Accent 1'
    sw_headers = ["Software Component", "Details"]
    sw_rows = [
        ("Front-End", "HTML, CSS, JavaScript (React JS framework, Vite)"),
        ("Back-End", "Java (Spring Boot framework)"),
        ("Database", "NoSQL (MongoDB)"),
        ("Web Server", "Embedded Apache Tomcat (shipped with Spring Boot)"),
        ("Operating System", "Windows 10 / 11")
    ]
    
    # Write Headers
    hdr_cells = sw_table.rows[0].cells
    hdr_cells[0].text = sw_headers[0]
    hdr_cells[1].text = sw_headers[1]
    for cell in hdr_cells:
        set_cell_background(cell, "7C4DFF")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell)
        
    # Write Rows
    for idx, (comp, det) in enumerate(sw_rows):
        row_cells = sw_table.rows[idx+1].cells
        row_cells[0].text = comp
        row_cells[1].text = det
        for cell in row_cells:
            set_cell_background(cell, "F3E5F5" if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    h3_2_1 = doc.add_paragraph()
    h3_2_1_run = h3_2_1.add_run("3.2.1. ABOUT THE SOFTWARE (FRONT END)")
    h3_2_1_run.font.size = Pt(12)
    h3_2_1_run.bold = True
    h3_2_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    r = p.add_run("An Introduction to React JS\n")
    r.bold = True
    p.add_run(
        "React JS is an open-source, component-based frontend JavaScript library developed by Meta. "
        "It provides developers with the capability to quickly build fast and responsive Single Page "
        "Web Applications (SPAs). React manages state changes inside user interfaces efficiently by "
        "rendering components only when data is updated, avoiding costly complete DOM tree re-renderings. "
        "Several modern companies (Netflix, Airbnb, WhatsApp) utilize React for UI development. "
        "React operates seamlessly alongside stylesheets, Lucide SVG icons, and HTML markup, ensuring "
        "dynamic user dashboards are rendered beautifully."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Characteristics of React JS\n")
    r.bold = True
    p.add_run(
        "React revolves around component-based modular structure and virtual DOM trees. "
        "Five key characteristics make React the tool of choice for frontend developers:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    react_features = [
        "Familiarity (uses standard JavaScript XML syntax / JSX)",
        "Component Reusability (UI is split into modular buttons, cards, inputs)",
        "Efficiency (Virtual DOM updates state quickly without freezing page)",
        "Security (JSX automatically sanitizes inputs to prevent Cross-Site Scripting)",
        "State Management (React hooks like useState and useEffect track variables easily)"
    ]
    for feat in react_features:
        doc.add_paragraph(feat, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    h3_2_2 = doc.add_paragraph()
    h3_2_2_run = h3_2_2.add_run("3.2.2. ABOUT THE DATABASE (BACK END)")
    h3_2_2_run.font.size = Pt(12)
    h3_2_2_run.bold = True
    h3_2_2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    r = p.add_run("Introduction to MongoDB\n")
    r.bold = True
    p.add_run(
        "MongoDB is a robust, open-source, document-based NoSQL database server. Documents are stored in "
        "JSON-like formats (BSON) containing fields and values. MongoDB eliminates traditional rigid SQL tables, "
        "enabling flexible schemas that map directly to Java objects. This makes MongoDB incredibly fast for reading "
        "and writing travel items. MongoDB Atlas is the cloud version, providing secure whitelisted IP clustering, "
        "scalability, and global online database availability."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Installation & Configuration of MongoDB\n")
    r.bold = True
    p.add_run(
        "MongoDB can be set up locally as a background Windows Service on port 27017, or deployed as a cloud cluster "
        "on MongoDB Atlas. Configuration requires establishing connection URIs. In Spring Boot, this is managed via "
        "the application.properties file by setting the 'spring.data.mongodb.uri' attribute to target either localhost "
        "or your Atlas SRV endpoint. Once connected, Spring Boot automatically instantiates repository beans to query "
        "database documents."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h3_2_3 = doc.add_paragraph()
    h3_2_3_run = h3_2_3.add_run("3.2.3. INTRODUCTION TO JAVASCRIPT")
    h3_2_3_run.font.size = Pt(12)
    h3_2_3_run.bold = True
    h3_2_3.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "JavaScript brings professional programming techniques and dynamic scripting to web browsers. "
        "With JavaScript, we can handle frontend login validation, manipulate elements, trigger toast notifications, "
        "and fetch backend API data asynchronously using the browser Fetch API. It brings true client-side processing "
        "and event handling, ensuring UI transitions feel immediate and alive."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 4 COVER ====================
    add_chapter_cover("SYSTEM DESIGN")
    
    # ==================== CHAPTER 4 CONTENT ====================
    doc.add_page_break()
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. SYSTEM DESIGN")
    h4_run.font.size = Pt(16)
    h4_run.bold = True
    h4.paragraph_format.space_after = Pt(12)
    
    h4_1 = doc.add_paragraph()
    h4_1_run = h4_1.add_run("4.1. DATA FLOW DIAGRAM")
    h4_1_run.font.size = Pt(14)
    h4_1_run.bold = True
    h4_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "A Data Flow Diagram (DFD) represents the flow of data through an information system. "
        "It depicts how input data is processed through services to update database stores. "
        "The standard symbols used in DFDs include:"
    )
    p.paragraph_format.space_after = Pt(6)
    
    dfd_symbols = [
        "Square / Rectangle: Represents a data source or destination (external entity like User/Admin).",
        "Open-Ended Rectangle: Represents data storage (database collections like users, packing_items).",
        "Directed Arrow: Represents flow of data.",
        "Oval / Circle: Represents a process transforming data streams (Login, Create Item, Remove User)."
    ]
    for sym in dfd_symbols:
        doc.add_paragraph(sym, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    h4_2 = doc.add_paragraph()
    h4_2_run = h4_2.add_run("4.2. SYSTEM ARCHITECTURE DIAGRAM")
    h4_2_run.font.size = Pt(14)
    h4_2_run.bold = True
    h4_2.paragraph_format.space_after = Pt(6)
    
    # Text-based Architecture flow
    p = doc.add_paragraph()
    r1 = p.add_run("CLIENT FLOW:\n")
    r1.bold = True
    p.add_run("User -> Login & Register -> Dashboard View -> Create/Modify Items -> Updates `packing_items` collection.\n\n")
    r2 = p.add_run("ADMIN FLOW:\n")
    r2.bold = True
    p.add_run("Admin -> Login -> Admin Dashboard -> View Directory / Delete User -> Wipes `users` & triggers cascading wipe on `packing_items`.")
    p.paragraph_format.space_after = Pt(12)
    
    h4_3 = doc.add_paragraph()
    h4_3_run = h4_3.add_run("4.3. DATABASE DESIGN")
    h4_3_run.font.size = Pt(14)
    h4_3_run.bold = True
    h4_3.paragraph_format.space_after = Pt(6)
    
    h4_3_1 = doc.add_paragraph()
    h4_3_1_run = h4_3_1.add_run("4.3.1. COLLECTION SCHEMA DESIGN")
    h4_3_1_run.font.size = Pt(12)
    h4_3_1_run.bold = True
    h4_3_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    r = p.add_run("Collection: users\n")
    r.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    # Users Collection Table
    u_table = doc.add_table(rows=6, cols=3)
    u_table.style = 'Light Shading Accent 1'
    u_headers = ["Field name", "Type", "Constraints"]
    u_rows = [
        ("id", "String", "PRIMARY KEY (Auto-generated)"),
        ("username", "String", "NOT NULL, UNIQUE"),
        ("email", "String", "NOT NULL, UNIQUE"),
        ("password", "String", "NOT NULL"),
        ("role", "String", "NOT NULL (CLIENT / ADMIN)")
    ]
    
    # Write User Table Headers
    hdr_cells = u_table.rows[0].cells
    for col_idx, text in enumerate(u_headers):
        hdr_cells[col_idx].text = text
    for cell in hdr_cells:
        set_cell_background(cell, "7C4DFF")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell)
        
    # Write User Table Rows
    for idx, (f, t, c) in enumerate(u_rows):
        row_cells = u_table.rows[idx+1].cells
        row_cells[0].text = f
        row_cells[1].text = t
        row_cells[2].text = c
        for cell in row_cells:
            set_cell_background(cell, "F3E5F5" if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Collection: packing_items\n")
    r.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    # Items Table
    i_table = doc.add_table(rows=6, cols=3)
    i_table.style = 'Light Shading Accent 1'
    i_headers = ["Field name", "Type", "Constraints"]
    i_rows = [
        ("id", "String", "PRIMARY KEY (Auto-generated)"),
        ("name", "String", "NOT NULL"),
        ("quantity", "Integer", "NOT NULL, DEFAULT 1"),
        ("packed", "Boolean", "NOT NULL, DEFAULT false"),
        ("userId", "String", "NOT NULL, FOREIGN KEY (References users.id)")
    ]
    
    # Write Items Table Headers
    hdr_cells = i_table.rows[0].cells
    for col_idx, text in enumerate(i_headers):
        hdr_cells[col_idx].text = text
    for cell in hdr_cells:
        set_cell_background(cell, "7C4DFF")
        cell.paragraphs[0].runs[0].font.bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        set_cell_margins(cell)
        
    # Write Items Table Rows
    for idx, (f, t, c) in enumerate(i_rows):
        row_cells = i_table.rows[idx+1].cells
        row_cells[0].text = f
        row_cells[1].text = t
        row_cells[2].text = c
        for cell in row_cells:
            set_cell_background(cell, "F3E5F5" if idx % 2 == 0 else "FFFFFF")
            set_cell_margins(cell)
            
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 5 COVER ====================
    add_chapter_cover("SYSTEM TESTING")
    
    # ==================== CHAPTER 5 CONTENT ====================
    doc.add_page_break()
    h5 = doc.add_paragraph()
    h5_run = h5.add_run("5. SYSTEM TESTING")
    h5_run.font.size = Pt(16)
    h5_run.bold = True
    h5.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph(
        "System testing is a crucial phase in the development of the travel packing list organizer. "
        "It involves testing the complete, integrated frontend web client and Java backend system to "
        "ensure all CRUD functionalities, CORS permissions, role checks, and database cascading operations "
        "work correctly according to requirements. This helps in identifying connection timeouts, preflight blocks, "
        "and data integrity bugs before deployment. Proper system testing ensures high reliability, operational efficiency, "
        "and a smooth user experience."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("5.1 Functional Testing\n")
    r.bold = True
    p.add_run(
        "Functional testing is performed to verify that all features, such as user login, password matching, "
        "item creation, packed-state checkboxes, and admin dashboard account deletes produce correct outputs for given inputs."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("5.2 Usability Testing\n")
    r.bold = True
    p.add_run(
        "Usability testing focuses on evaluating UI friendliness. It ensures users can add items and click checkboxes "
        "easily, and stats calculate instantly. It checks that modals look clear on both mobile screen widths and desktop monitors."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("5.3 Performance Testing\n")
    r.bold = True
    p.add_run(
        "Performance testing checks how well the system performs under different loads, such as multiple users "
        "accessing lists at the same time. MongoDB NoSQL databases are tested to ensure read/write queries return quickly."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("5.4 Security Testing\n")
    r.bold = True
    p.add_run(
        "Security testing ensures user directories are protected. It verifies that non-authenticated users are redirected "
        "back to `/login` via React Route guards and that clients cannot access admin dashboard pages."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("5.5 Database Testing\n")
    r.bold = True
    p.add_run(
        "Database testing verifies CRUD integrity in MongoDB Atlas. It confirms that the delete user trigger deletes all "
        "referenced items correctly, leaving no orphaned data in the collection."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 6 COVER ====================
    add_chapter_cover("SYSTEM IMPLEMENTATION")
    
    # ==================== CHAPTER 6 CONTENT ====================
    doc.add_page_break()
    h6 = doc.add_paragraph()
    h6_run = h6.add_run("6. SYSTEM IMPLEMENTATION")
    h6_run.font.size = Pt(16)
    h6_run.bold = True
    h6.paragraph_format.space_after = Pt(12)
    
    h6_1 = doc.add_paragraph()
    h6_1_run = h6_1.add_run("6.1. CODING")
    h6_1_run.font.size = Pt(14)
    h6_1_run.bold = True
    h6_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph()
    r = p.add_run("config.js (API Base Config)\n")
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    
    # Code block style helper
    def add_code_block(code_text):
        p_code = doc.add_paragraph()
        p_code.paragraph_format.left_indent = Inches(0.2)
        p_code.paragraph_format.space_after = Pt(12)
        run_code = p_code.add_run(code_text)
        run_code.font.name = 'Courier New'
        run_code.font.size = Pt(10)
        
    add_code_block(
        "// Automatically connects to localhost when developing, and to Render when live\n"
        "export const API_BASE =\n"
        "  window.location.hostname === 'localhost' || window.location.hostname === '127.0.0.1'\n"
        "    ? 'http://localhost:8080'\n"
        "    : 'https://packing-list-backend.onrender.com';"
    )
    
    p = doc.add_paragraph()
    r = p.add_run("PackingListApplication.java (CORS & Admin Seed)\n")
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(
        "@SpringBootApplication\n"
        "public class PackingListApplication {\n"
        "    public static void main(String[] args) {\n"
        "        SpringApplication.run(PackingListApplication.class, args);\n"
        "    }\n\n"
        "    // Seed default admin user on startup if not present\n"
        "    @Bean\n"
        "    public CommandLineRunner initDatabase(UserRepository userRepository) {\n"
        "        return args -> {\n"
        "            if (userRepository.findByUsername(\"admin\").isEmpty()) {\n"
        "                User admin = new User();\n"
        "                admin.setUsername(\"admin\");\n"
        "                admin.setPassword(\"admin123\");\n"
        "                admin.setEmail(\"admin@packmate.com\");\n"
        "                admin.setRole(\"ADMIN\");\n"
        "                userRepository.save(admin);\n"
        "            }\n"
        "        };\n"
        "    }\n"
        "}"
    )
    
    p = doc.add_paragraph()
    r = p.add_run("AuthController.java (Secure Router Endpoint)\n")
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(
        "@CrossOrigin(originPatterns = {\"http://localhost:5173\", \"https://*.vercel.app\"}, allowCredentials = \"true\")\n"
        "@RestController\n"
        "@RequestMapping(\"/api/auth\")\n"
        "public class AuthController {\n"
        "    @Autowired\n"
        "    private UserService userService;\n\n"
        "    @PostMapping(\"/login\")\n"
        "    public ResponseEntity<?> login(@RequestBody LoginRequest loginRequest) {\n"
        "        Optional<User> userOpt = userService.login(loginRequest.getUsername(), loginRequest.getPassword());\n"
        "        if (userOpt.isPresent()) {\n"
        "            User user = userOpt.get();\n"
        "            Map<String, Object> response = new HashMap<>();\n"
        "            response.put(\"id\", user.getId());\n"
        "            response.put(\"username\", user.getUsername());\n"
        "            response.put(\"role\", user.getRole());\n"
        "            return ResponseEntity.ok(response);\n"
        "        } else {\n"
        "            return ResponseEntity.status(HttpStatus.UNAUTHORIZED).body(\"Invalid credentials!\");\n"
        "        }\n"
        "    }\n"
        "}"
    )
    
    p = doc.add_paragraph()
    r = p.add_run("logout logic (App.jsx snippet)\n")
    r.italic = True
    p.paragraph_format.space_after = Pt(6)
    
    add_code_block(
        "const handleLogout = () => {\n"
        "  triggerToast('Logged out successfully.');\n"
        "  setUser(null);\n"
        "  localStorage.removeItem('packing_list_user');\n"
        "};"
    )
    
    h6_2 = doc.add_paragraph()
    h6_2_run = h6_2.add_run("6.2. SCREEN LAYOUT")
    h6_2_run.font.size = Pt(14)
    h6_2_run.bold = True
    h6_2.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "The screen layouts are fully responsive and styled with custom purple-black glassmorphic panels. "
        "You can refer to the screenshots section in the college project document directory. It contains detailed "
        "wireframes for the Login Page, User Registration Page, Client Checklist Dashboard, and Admin Moderation Dashboard."
    )
    p.paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 7 COVER ====================
    add_chapter_cover("CONCLUSION & FUTURE ENHANCEMENTS")
    
    # ==================== CHAPTER 7 CONTENT ====================
    doc.add_page_break()
    h7 = doc.add_paragraph()
    h7_run = h7.add_run("7. CONCLUSION AND FUTURE ENHANCEMENTS")
    h7_run.font.size = Pt(16)
    h7_run.bold = True
    h7.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph(
        "This section details the conclusions drawn from system execution and potential enhancements to extend the travel organizer in the future."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h7_1 = doc.add_paragraph()
    h7_1_run = h7_1.add_run("7.1. CONCLUSION")
    h7_1_run.font.size = Pt(14)
    h7_1_run.bold = True
    h7_1.paragraph_format.space_after = Pt(6)
    
    p = doc.add_paragraph(
        "The travel packing organizer application (PackMate) has been successfully designed, coded, and deployed. "
        "It provides secure credentials authentication, responsive user checklist CRUD capabilities, and direct "
        "admin audit dashboard features. The database cascade deletes all associated packing items when an account is removed, "
        "ensuring data consistency. Overall, it replaces paper checklists with a reliable, dynamic, and modern tool."
    )
    p.paragraph_format.space_after = Pt(12)
    
    h7_2 = doc.add_paragraph()
    h7_2_run = h7_2.add_run("7.2. FUTURE ENHANCEMENTS")
    h7_2_run.font.size = Pt(14)
    h7_2_run.bold = True
    h7_2.paragraph_format.space_after = Pt(6)
    
    enhancements = [
        "Mobile Application Integration - Developing native iOS/Android applications using React Native for off-line synchronization.",
        "Automated Packing Suggestions - Implementing AI algorithms to recommend items based on travel location, climate, and length of stay.",
        "Group Packing - Collaborative real-time checklists allowing multiple travelers to coordinate luggage packing synchronously.",
        "Social Sharing - Exporting and sharing list templates with friends and travel groups.",
        "Data Encryption - Implementing secure BCrypt encryption inside the backend database service.",
        "Multi-Language Support - Adding translation strings to make the platform accessible to global travelers."
    ]
    for idx, enh in enumerate(enhancements):
        doc.add_paragraph(f"{idx+1}. {enh}")
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # ==================== CHAPTER 8 COVER ====================
    add_chapter_cover("BIBLIOGRAPHY")
    
    # ==================== CHAPTER 8 CONTENT ====================
    doc.add_page_break()
    h8 = doc.add_paragraph()
    h8_run = h8.add_run("8. BIBLIOGRAPHY AND REFERENCES")
    h8_run.font.size = Pt(16)
    h8_run.bold = True
    h8.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph(
        "Bibliography and references detail the textbooks, documentations, and web resources used in research and development."
    )
    p.paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Book References:\n")
    r.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    books = [
        "Spring Boot in Action, Craig Walls, Manning Publications, 2016.",
        "React Key Concepts, Maximilian Schwarzmuller, Packt Publishing, 2022.",
        "MongoDB: The Definitive Guide, Shannon Bradshaw, O'Reilly Media, 2019.",
        "JavaScript: The Definitive Guide, David Flanagan, O'Reilly Media, 2020."
    ]
    for b in books:
        doc.add_paragraph(b, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    p = doc.add_paragraph()
    r = p.add_run("Web References:\n")
    r.bold = True
    p.paragraph_format.space_after = Pt(6)
    
    webs = [
        "https://react.dev/ - React Official Documentation",
        "https://spring.io/projects/spring-boot - Spring Boot Guides",
        "https://www.mongodb.com/docs/ - MongoDB Official Database documentation",
        "https://vitejs.dev/ - Vite Bundler Documentation",
        "https://lucide.dev/ - Lucide SVG Icons Library"
    ]
    for w in webs:
        doc.add_paragraph(w, style='List Bullet')
        
    doc.add_paragraph().paragraph_format.space_after = Pt(12)
    
    # Save the document
    doc.save("Project_Report.docx")
    print("Project_Report.docx created successfully!")

if __name__ == "__main__":
    create_report()
